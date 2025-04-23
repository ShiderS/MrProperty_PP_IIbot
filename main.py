import asyncio
import os

import base64
from openai import AsyncOpenAI
from aiogram import Bot, Dispatcher, F, types
from aiogram.filters import CommandStart
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext
from aiogram.types import Message, FSInputFile, BufferedInputFile
from docx import Document

from config import TG_TOKEN, OPENAI_API_KEY, DB_NAME, TEMP_IMAGE_FOLDER, TEMP_DOC_FOLDER, OPENAI_VISION_MODEL
from data import db_session
from data.user import User

bot = Bot(token=TG_TOKEN)
dp = Dispatcher()

openai_client = AsyncOpenAI(api_key=OPENAI_API_KEY)


class ImageProcessing(StatesGroup):
    waiting_for_images = State()

def ensure_dir_exists(dir_path):
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)

# async def get_qwen_text_from_image(image_path: str):
#     messages = [{
#         "role": "user",
#         "content": [
#             {"image": f"file://{os.path.abspath(image_path)}"},
#             {"text": "Извлеки весь текст с этого изображения конспекта. Сохрани форматирование, насколько это возможно."}
#         ]
#     }]
#     try:
#         logging.info(f"Отправка запроса в Qwen API для файла: {image_path}")
#         response = await dashscope.MultiModalConversation.aCall(
#             model='qwen-vl-plus',
#             messages=messages
#         )
#
#         if response.status_code == HTTPStatus.OK:
#             logging.info(f"Qwen API ответил успешно для файла: {image_path}")
#             content = response.output.choices[0].message.content
#             if isinstance(content, list) and len(content) > 0 and 'text' in content[0]:
#                 return content[0]['text'].strip()
#             else:
#                 logging.warning(f"Неожиданная структура ответа от Qwen: {response}")
#                 return None
#         else:
#             logging.error(f'Ошибка Qwen API: Код={response.code}, Сообщение={response.message}, RequestId={response.request_id}')
#             return None
#
#     except Exception as e:
#         logging.exception(f"Исключение при вызове Qwen API для {image_path}: {e}")
#         return None


# --- Обработчики команд и сообщений ---
def image_to_base64(image_path: str):
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    except Exception as e:
        return None

async def get_openai_text_from_image(image_path: str):
    base64_image = image_to_base64(image_path)
    if not base64_image:
        return None

    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "input_text", "text": "Извлеки весь текст с этого изображения рукописного конспекта. Постарайся сохранить структуру, абзацы и переносы строк, как в оригинале."
                },
                {
                    "type": "input_image",
                    "image_url": f"data:image/jpeg;base64,{base64_image}"
                }
            ]
        }
    ]

    try:
        response = await openai_client.chat.completions.create(
            model=OPENAI_VISION_MODEL,
            messages=messages,
        )


        if response.choices and response.choices[0].message and response.choices[0].message.content:
            extracted_text = response.choices[0].message.content.strip()
            return extracted_text
        else:
            return None
    except Exception as e:
        return "[Произошла внутренняя ошибка при обработке изображения]"


def create_word_document(text: str, filename: str):
    try:
        document = Document()
        for paragraph_text in text.split('\n'):
             if paragraph_text.strip():
                 document.add_paragraph(paragraph_text)
             else:
                 document.add_paragraph()

        ensure_dir_exists(TEMP_DOC_FOLDER)
        doc_path = os.path.join(TEMP_DOC_FOLDER, filename)
        document.save(doc_path)
        return doc_path
    except Exception as e:
        return None

@dp.message(CommandStart())
async def cmd_start(message: Message, state: FSMContext):
    await state.clear()
    session = db_session.create_session()
    try:
        user = session.query(User).filter(User.id == message.from_user.id).first()
        if not user:
            user = User(
                id=message.from_user.id,
                full_name=message.from_user.full_name,
                tg_name=message.from_user.username
            )
            session.add(user)
            session.commit()
        else:
            await message.answer(
                f"Привет, {message.from_user.first_name}! 👋\n\n"
                "Отправь мне одну или несколько картинок твоего конспекта. "
                "Когда закончишь, нажми кнопку 'Готово ✅' или отправь команду /done, "
                "и я переведу их в текст и пришлю Word-файл."
        )
        await state.set_state(ImageProcessing.waiting_for_images)
        await state.update_data(image_files=[])
    except:
        pass

@dp.message(ImageProcessing.waiting_for_images, F.photo)
async def handle_photos(message: Message, state: FSMContext):
    if not message.photo:
        await message.reply("Пожалуйста, отправьте изображение.")
        return

    photo = message.photo[-1]
    file_id = photo.file_id

    user_data = await state.get_data()
    image_files = user_data.get("image_files", [])

    image_files.append(file_id)

    await state.update_data(image_files=image_files)


    keyboard = types.InlineKeyboardMarkup(inline_keyboard=[
        [types.InlineKeyboardButton(text="Готово ✅", callback_data="process_images")]
    ])


    try:
        await message.reply(
        f"Фото {len(image_files)} принято! 👍 Отправляй еще или нажми 'Готово ✅'.",
        reply_markup=keyboard
        )
    except Exception as e:
              await message.reply(
                 f"Фото {len(image_files)} принято! 👍 Отправляй еще или нажми 'Готово ✅'.",
                 reply_markup=keyboard
             )


@dp.message(ImageProcessing.waiting_for_images, F.text == "/done")
async def handle_done_command(message: Message, state: FSMContext):
    await process_uploaded_images(message, state)


@dp.callback_query(ImageProcessing.waiting_for_images, F.data == "process_images")
async def handle_done_button(callback: types.CallbackQuery, state: FSMContext):
    await callback.answer("Начинаю обработку...")
    await process_uploaded_images(callback.message, state, is_callback=True)


async def process_uploaded_images(message_or_callback_message: Message, state: FSMContext, is_callback: bool = False):
    user_data = await state.get_data()
    image_files = user_data.get("image_files", [])

    if not image_files:
        await message_or_callback_message.answer("Ты не отправил ни одного изображения. Отправь фото и потом нажми 'Готово ✅'.")
        return

    await state.clear()

    processing_message = await message_or_callback_message.answer(f"Получено {len(image_files)} фото. Начинаю распознавание текста... 🧠")

    all_extracted_text = []
    temp_files_to_delete = []
    ensure_dir_exists(TEMP_IMAGE_FOLDER)

    for i, file_id in enumerate(image_files):
        try:
            await processing_message.edit_text(f"Обрабатываю фото {i+1} из {len(image_files)}...")
            file_info = await bot.get_file(file_id)
            file_path = file_info.file_path
            temp_image_path = os.path.join(TEMP_IMAGE_FOLDER, f"{message_or_callback_message.chat.id}_{file_id}.jpeg")
            temp_files_to_delete.append(temp_image_path)

            await bot.download_file(file_path, temp_image_path)

            extracted_text = await get_openai_text_from_image(temp_image_path)

            if extracted_text:
                all_extracted_text.append(extracted_text)
            else:
                all_extracted_text.append(f"[Не удалось распознать текст на изображении {i+1}]")
        except Exception as e:
            all_extracted_text.append(f"[Ошибка при обработке изображения {i+1}]")
        finally:
            pass

    await processing_message.edit_text("Текст извлечен. Создаю Word документ... 📄")
    final_text = "\n\n---\n\n".join(all_extracted_text)

    doc_filename = f"konspekt_{message_or_callback_message.chat.id}_{message_or_callback_message.message_id}.docx"
    doc_path = create_word_document(final_text, doc_filename)

    if doc_path:
        try:
            input_file = FSInputFile(doc_path, filename=f"Твой_конспект_{message_or_callback_message.chat.id}.docx")
            await message_or_callback_message.answer_document(input_file, caption="Готово! Вот твой конспект в формате Word.")
            try:
                os.remove(doc_path)
            except OSError as e:
                pass

        except Exception as e:
            await message_or_callback_message.answer("Не смог отправить Word файл. Попробуй еще раз позже.")
        finally:
             await processing_message.delete() # Удаляем сообщение "Создаю Word..."

    else:
        await processing_message.edit_text("Не удалось создать Word файл. Отправляю текст как сообщение:")
        max_length = 4096
        for i in range(0, len(final_text), max_length):
            await message_or_callback_message.answer(final_text[i:i + max_length])

    for f_path in temp_files_to_delete:
        try:
            os.remove(f_path)
        except OSError as e:
            pass

@dp.message(ImageProcessing.waiting_for_images)
async def handle_other_messages_while_waiting(message: Message, state: FSMContext):
    await message.reply("Пожалуйста, отправь мне фото конспекта или нажми 'Готово ✅' / отправь /done, если закончил.")


@dp.message()
async def handle_unknown_messages(message: Message, state: FSMContext):
    current_state = await state.get_state()
    if current_state is None:
        await message.reply("Я не понимаю эту команду. Используй /start, чтобы начать.")


# --- Запуск бота ---
async def main():
    db_session.global_init(DB_NAME)
    ensure_dir_exists(TEMP_IMAGE_FOLDER)
    ensure_dir_exists(TEMP_DOC_FOLDER)

    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())